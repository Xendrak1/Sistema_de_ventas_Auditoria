from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from build_documentacion import (
    OUT,
    PROJECT_TITLE,
    SOURCES,
    add_bullet,
    add_caption,
    add_cover,
    add_docx_evidence_image,
    add_number,
    add_paragraph,
    polish_document,
    set_cell_text,
    set_doc_style,
    style_table,
)


def add_small_note(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_text(cell, f"{title}: {body}")
    style_table(table, header_fill="F4F6F9")


def add_image_with_context(doc, filename, caption, context, what_to_show, control):
    add_docx_evidence_image(doc, filename, caption, context, width=5.7)
    table = doc.add_table(rows=1, cols=2)
    set_cell_text(table.rows[0].cells[0], "Qué demuestra", True)
    set_cell_text(table.rows[0].cells[1], what_to_show, False)
    row = table.add_row().cells
    set_cell_text(row[0], "Control asociado", True)
    set_cell_text(row[1], control, False)
    style_table(table)


def build_password_policy():
    doc = Document()
    set_doc_style(doc)
    add_cover(doc, "Política de uso de contraseñas", PROJECT_TITLE)

    doc.add_heading("1. Propósito", level=1)
    add_paragraph(
        doc,
        "Esta política define las reglas mínimas para crear, usar, proteger, recuperar y administrar contraseñas en el sistema web de ventas y en el servidor Fedora 39 utilizado para la auditoría. Su objetivo es reducir accesos no autorizados y evitar que una contraseña débil comprometa el sistema completo.",
    )

    doc.add_heading("2. Alcance", level=1)
    add_paragraph(
        doc,
        "La política aplica a los usuarios del sistema de ventas, al administrador del servidor Fedora, al equipo de desarrollo y a cualquier persona que tenga acceso al repositorio, a la base de datos, a la máquina virtual o al entorno de demostración.",
    )

    doc.add_heading("3. Principios de la política", level=1)
    for item in [
        "Cada usuario debe tener una cuenta individual. No se recomienda compartir usuarios genéricos salvo en demostraciones controladas.",
        "La contraseña debe proteger el acceso al sistema y no debe escribirse en documentos, chats o capturas destinadas a publicación.",
        "Las credenciales usadas para la práctica, como admin/admin123, existen solo para demostrar una falla prediseñada y no deben utilizarse en producción.",
        "El almacenamiento seguro de contraseñas es responsabilidad del sistema; el usuario no debería depender únicamente de su memoria o cuidado personal.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("4. Reglas para creación de contraseñas", level=1)
    table = doc.add_table(rows=1, cols=3)
    for i, header in enumerate(["Regla", "Requisito", "Motivo"]):
        set_cell_text(table.rows[0].cells[i], header, True)
    rows = [
        ["Longitud", "Mínimo 10 caracteres", "A mayor longitud, mayor resistencia ante ataques de adivinación."],
        ["Complejidad", "Combinar letras, números y símbolos", "Dificulta contraseñas demasiado simples o predecibles."],
        ["Evitar datos personales", "No usar nombre, CI, fecha de nacimiento o teléfono", "Estos datos pueden ser conocidos por otras personas."],
        ["Evitar contraseñas comunes", "No usar admin123, 123456, password, ventas123", "Son las primeras opciones probadas en ataques básicos."],
        ["No reutilizar", "No repetir la misma contraseña en otros sistemas", "Evita que una filtración externa afecte este sistema."],
    ]
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            set_cell_text(cells[i], value)
    style_table(table)

    doc.add_heading("5. Controles técnicos recomendados", level=1)
    controls = [
        ["Hash de contraseñas", "Guardar contraseñas con bcrypt o argon2, nunca en texto plano."],
        ["Bloqueo por intentos fallidos", "Bloquear temporalmente la cuenta después de 5 intentos fallidos."],
        ["Mensajes de error discretos", "No indicar si falló el usuario o la contraseña; solo mostrar credenciales incorrectas."],
        ["Recuperación segura", "Restablecer contraseñas mediante un proceso controlado, no mostrando la contraseña actual."],
        ["Registro de eventos", "Registrar intentos fallidos, cambios de contraseña y accesos administrativos."],
        ["Sesiones seguras", "Configurar cookies HttpOnly, SameSite y Secure cuando se use HTTPS."],
    ]
    table = doc.add_table(rows=1, cols=2)
    set_cell_text(table.rows[0].cells[0], "Control", True)
    set_cell_text(table.rows[0].cells[1], "Aplicación", True)
    for row_data in controls:
        cells = table.add_row().cells
        set_cell_text(cells[0], row_data[0])
        set_cell_text(cells[1], row_data[1])
    style_table(table)

    doc.add_heading("6. Responsabilidades", level=1)
    for item in [
        "Usuario: mantener en reserva sus credenciales y cerrar sesión al finalizar.",
        "Administrador: crear cuentas, revocar accesos, revisar intentos fallidos y proteger el servidor.",
        "Desarrollador: implementar hash, validaciones y controles de sesión.",
        "Equipo auditor: verificar que las reglas propuestas mitiguen los riesgos encontrados.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("7. Excepción para el laboratorio", level=1)
    add_paragraph(
        doc,
        "Durante la auditoría se permiten credenciales débiles documentadas para demostrar el riesgo. Esta excepción solo aplica dentro del laboratorio de VirtualBox y durante la presentación. Al finalizar, el servicio debe detenerse o la máquina virtual debe apagarse.",
    )

    doc.add_heading("8. Conclusión", level=1)
    add_paragraph(
        doc,
        "Una política de contraseñas no solo define cómo debe escribir una clave el usuario; también obliga al sistema a protegerla correctamente. En este proyecto se evidenció que una contraseña débil y almacenada en texto plano puede comprometer por completo el acceso al sistema. Por ello, antes de usar una aplicación similar en un entorno real, se deben aplicar controles de longitud, bloqueo, registro de eventos y almacenamiento seguro con hash.",
    )

    doc.add_heading("Bibliografía", level=1)
    for source in [SOURCES[3], SOURCES[5]]:
        add_paragraph(doc, source)

    polish_document(doc)
    output = OUT / "Politica_Uso_Contrasenas.docx"
    doc.save(output)
    return output


def build_secure_use_manual():
    doc = Document()
    set_doc_style(doc)
    add_cover(doc, "Manual de uso seguro y establecimiento de controles", PROJECT_TITLE)

    doc.add_heading("1. Finalidad del manual", level=1)
    add_paragraph(
        doc,
        "Este documento funciona como el documento complementario de establecimiento de controles solicitado para el proyecto. Está escrito como un manual práctico para dos públicos: el usuario que opera el sistema web de ventas y el administrador que prepara el servidor Fedora durante la demostración. También sirve como guía breve para explicar qué se hizo en la auditoría y qué controles se deberían aplicar si el sistema pasara a un entorno real.",
    )
    add_small_note(
        doc,
        "Uso del documento",
        "No reemplaza al informe principal; lo complementa con reglas claras, pasos de uso seguro, controles de contraseñas y evidencia visual del sistema y servidor.",
    )

    doc.add_heading("2. Roles dentro del laboratorio", level=1)
    table = doc.add_table(rows=1, cols=3)
    for i, header in enumerate(["Rol", "Responsabilidad", "Ejemplo en el proyecto"]):
        set_cell_text(table.rows[0].cells[i], header, True)
    rows = [
        ["Usuario del sistema", "Ingresar al sistema, registrar datos y cerrar sesión correctamente.", "Uso de login, productos, clientes, ventas y reportes."],
        ["Administrador del servidor", "Levantar Fedora, ejecutar Flask, abrir/cerrar puerto y revisar servicios.", "Servidor Fedora 39 con IP 192.168.56.103."],
        ["Auditor", "Probar acceso, ejecutar Nmap/Nessus y validar fallas prediseñadas.", "Kali Linux con IP 192.168.56.102."],
        ["Desarrollador", "Corregir vulnerabilidades y aplicar controles.", "Código fuente en GitHub."],
    ]
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            set_cell_text(cells[i], value)
    style_table(table)

    doc.add_heading("3. Preparación del servidor Fedora", level=1)
    add_paragraph(
        doc,
        "Antes de usar el sistema, Fedora debe estar en la misma red que Kali. En el laboratorio se usó una red solo anfitrión de VirtualBox. Esto permite que Kali acceda a Fedora sin publicar el sistema en internet.",
    )
    add_image_with_context(
        doc,
        "evidencia_docx_03.png",
        "Figura 1. Direcciones IP del laboratorio.",
        "La imagen muestra las direcciones de red de las máquinas virtuales. Fedora utiliza 192.168.56.103 y Kali utiliza 192.168.56.102.",
        "Confirma que ambas máquinas están comunicadas dentro de la misma red de laboratorio.",
        "Mantener el laboratorio aislado y no exponer el sistema vulnerable a internet.",
    )
    add_paragraph(doc, "Pasos recomendados para iniciar el servidor:")
    for step in [
        "Encender la máquina virtual Fedora 39.",
        "Abrir una terminal dentro de Fedora.",
        "Entrar a la carpeta del proyecto Sistema_de_ventas_Auditoria.",
        "Activar el entorno virtual de Python.",
        "Ejecutar python app.py.",
        "Verificar que el servicio escuche en el puerto 5000.",
    ]:
        add_number(doc, step)
    add_image_with_context(
        doc,
        "evidencia_docx_11.png",
        "Figura 2. Servicio Flask ejecutándose en Fedora.",
        "La captura muestra el sistema ejecutándose con Python/Flask en Fedora. También se observan solicitudes HTTP recibidas por el servidor.",
        "Demuestra que el sistema no es solo código fuente, sino que fue desplegado y ejecutado.",
        "Ejecutar el servicio solo durante la práctica y detenerlo al finalizar.",
    )

    doc.add_heading("4. Ingreso seguro al sistema web", level=1)
    add_paragraph(
        doc,
        "El usuario debe ingresar mediante la pantalla de login. Para la práctica se usaron credenciales débiles con fines académicos, pero en un sistema real estas credenciales no serían aceptables.",
    )
    add_image_with_context(
        doc,
        "evidencia_docx_05.png",
        "Figura 3. Pantalla de inicio de sesión.",
        "La pantalla de login se abre desde Kali, lo que demuestra que el sistema web de Fedora es accesible desde otra máquina.",
        "Permite evidenciar la autenticación del sistema y el punto donde se probaron credenciales débiles y SQL Injection.",
        "Aplicar contraseñas robustas, bloqueo por intentos fallidos y consultas parametrizadas.",
    )
    add_paragraph(doc, "Reglas para el usuario al iniciar sesión:")
    for item in [
        "No compartir usuario ni contraseña.",
        "No dejar la sesión abierta si otra persona usará el equipo.",
        "No guardar credenciales en capturas, chats o documentos públicos.",
        "Reportar cualquier acceso extraño o comportamiento inesperado.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("5. Uso de módulos principales", level=1)
    add_image_with_context(
        doc,
        "evidencia_docx_01.png",
        "Figura 4. Panel principal del sistema.",
        "El panel muestra productos, clientes, ventas e ingresos. También presenta alertas de inventario.",
        "Sirve como vista de control para que el usuario revise el estado general del sistema.",
        "Proteger el panel con autenticación y cerrar sesión después de usarlo.",
    )
    add_image_with_context(
        doc,
        "evidencia_docx_19.png",
        "Figura 5. Módulo de productos.",
        "El módulo de productos muestra el catálogo, imágenes, precios, categorías y stock.",
        "Es una sección importante porque procesa entradas como nombre y descripción del producto.",
        "Validar y sanitizar los campos para evitar XSS almacenado.",
    )
    add_image_with_context(
        doc,
        "evidencia_docx_06.png",
        "Figura 6. Módulo de clientes.",
        "La imagen muestra el listado y registro de clientes dentro del sistema.",
        "Contiene datos que deben tratarse con cuidado aunque sean de prueba.",
        "Limitar acceso, validar campos y evitar registrar información real en el laboratorio.",
    )
    add_image_with_context(
        doc,
        "evidencia_docx_02.png",
        "Figura 7. Búsqueda de productos.",
        "La búsqueda permite consultar productos por texto.",
        "Fue usada para demostrar XSS reflejado, ya que la entrada del usuario aparece en la respuesta.",
        "Escapar salidas HTML y validar parámetros de búsqueda.",
    )
    add_image_with_context(
        doc,
        "evidencia_docx_12.png",
        "Figura 8. Módulo de reportes.",
        "El módulo resume ventas por producto y el estado del inventario.",
        "Ayuda a explicar que el sistema tiene funcionalidad administrativa suficiente para ser auditado.",
        "Proteger reportes con autenticación y controlar quién puede consultarlos.",
    )

    doc.add_heading("6. Controles para el uso de contraseñas", level=1)
    table = doc.add_table(rows=1, cols=3)
    for i, header in enumerate(["Control", "Qué debe hacerse", "Por qué importa"]):
        set_cell_text(table.rows[0].cells[i], header, True)
    controls = [
        ["Contraseña robusta", "Usar mínimo 10 caracteres con combinación de letras, números y símbolos.", "Reduce el riesgo de adivinación."],
        ["Bloqueo por intentos", "Bloquear temporalmente después de varios intentos fallidos.", "Dificulta ataques de fuerza bruta."],
        ["Hash seguro", "Guardar contraseñas con bcrypt o argon2.", "Evita que la base revele claves reales."],
        ["No reutilización", "No usar la misma contraseña en otros sistemas.", "Evita que una filtración externa comprometa el sistema."],
        ["Cierre de sesión", "Usar la opción Salir al terminar.", "Evita accesos no autorizados desde una sesión abierta."],
    ]
    for row_data in controls:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            set_cell_text(cells[i], value)
    style_table(table)

    doc.add_heading("7. Controles para el servidor Fedora", level=1)
    for item in [
        "Abrir el puerto 5000 solo durante la demostración.",
        "Cerrar el servicio Flask al terminar la práctica.",
        "Revisar puertos activos con ss -tulpn.",
        "Mantener Fedora actualizado antes de la presentación.",
        "No ejecutar el sistema como root si no es necesario.",
        "Deshabilitar servicios innecesarios como LLMNR si no se usan.",
        "No publicar la IP del laboratorio como si fuera un servidor de producción.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("8. Guía breve para demostrar la auditoría", level=1)
    add_paragraph(doc, "Si el documento se usa durante la exposición, el orden recomendado para mostrar la práctica es el siguiente:")
    for step in [
        "Mostrar el repositorio de GitHub para evidenciar código fuente propio.",
        "Mostrar Fedora con el sistema Flask ejecutándose.",
        "Abrir el sistema desde Kali para evidenciar acceso desde otra máquina.",
        "Ejecutar o mostrar el resultado de Nmap.",
        "Mostrar el resultado de Nessus sin autenticación.",
        "Realizar las pruebas manuales: SQL Injection, XSS reflejado, XSS almacenado y contraseñas en texto plano.",
        "Relacionar cada hallazgo con la matriz de riesgos y los controles propuestos.",
    ]:
        add_number(doc, step)
    add_image_with_context(
        doc,
        "evidencia_docx_07.png",
        "Figura 9. Resultado de Nmap.",
        "Nmap detecta el puerto 5000/TCP abierto y reconoce el servicio Werkzeug/Python.",
        "Demuestra reconocimiento de servicios desde la máquina auditora.",
        "Controlar puertos expuestos y no usar el servidor de desarrollo Flask en producción.",
    )
    add_image_with_context(
        doc,
        "evidencia_docx_18.png",
        "Figura 10. Resultado de Nessus.",
        "Nessus muestra hallazgos informativos y de baja severidad sobre el servidor y el servicio HTTP.",
        "Demuestra que el servidor fue evaluado con una herramienta automática.",
        "Revisar cabeceras, cookies, exposición de información y servicios innecesarios.",
    )

    doc.add_heading("9. Pruebas manuales y controles asociados", level=1)
    table = doc.add_table(rows=1, cols=4)
    for i, header in enumerate(["Prueba", "Qué demuestra", "Riesgo", "Control recomendado"]):
        set_cell_text(table.rows[0].cells[i], header, True)
    rows = [
        ["SQL Injection", "El login acepta entradas maliciosas.", "Ingreso sin credenciales válidas.", "Usar consultas parametrizadas."],
        ["XSS reflejado", "La búsqueda refleja scripts.", "Ejecución de JavaScript en navegador.", "Escapar salidas HTML."],
        ["XSS almacenado", "La descripción guarda scripts.", "Código malicioso persistente.", "Sanitizar entradas antes de guardar."],
        ["Contraseñas en texto plano", "La base muestra claves legibles.", "Robo directo de credenciales.", "Aplicar hash seguro."],
    ]
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            set_cell_text(cells[i], value)
    style_table(table)
    add_image_with_context(
        doc,
        "evidencia_docx_10.png",
        "Figura 11. Prueba de SQL Injection.",
        "Se usa un payload en el login para demostrar que la consulta SQL no está protegida.",
        "Permite explicar una de las fallas más críticas del sistema.",
        "Reescribir consultas SQL usando parámetros.",
    )
    add_image_with_context(
        doc,
        "evidencia_docx_13.png",
        "Figura 12. Prueba de XSS reflejado.",
        "La alerta demuestra ejecución de JavaScript desde una entrada reflejada.",
        "Muestra por qué no se debe imprimir entrada del usuario sin escape.",
        "Escapar salidas y validar parámetros.",
    )
    add_image_with_context(
        doc,
        "evidencia_docx_21.png",
        "Figura 13. Prueba de XSS almacenado.",
        "La alerta confirma que el script quedó guardado y se ejecutó después.",
        "Tiene mayor impacto porque permanece en el sistema hasta corregir el registro.",
        "Sanitizar contenido y no usar datos del usuario como HTML seguro.",
    )
    add_image_with_context(
        doc,
        "evidencia_docx_08.png",
        "Figura 14. Contraseñas en texto plano.",
        "SQLite muestra usuarios y contraseñas legibles.",
        "Demuestra que una copia de la base compromete las cuentas.",
        "Guardar contraseñas con hash seguro y sal.",
    )

    doc.add_heading("10. Checklist final para el administrador", level=1)
    for item in [
        "Confirmar que Fedora y Kali estén en la misma red.",
        "Verificar que el sistema abra desde Kali.",
        "Tener capturas listas de Nmap, Nessus y pruebas manuales.",
        "Cerrar el puerto 5000 o apagar Fedora al finalizar.",
        "No reutilizar credenciales débiles fuera del laboratorio.",
        "Explicar que las fallas son prediseñadas y con fines académicos.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("11. Conclusión", level=1)
    add_paragraph(
        doc,
        "El manual permite convertir los hallazgos de la auditoría en reglas prácticas de uso. El sistema fue diseñado con fallas prediseñadas para aprender, pero eso no elimina la necesidad de establecer controles básicos. El usuario debe manejar correctamente sus credenciales y cerrar sesión; el administrador debe controlar el servidor, los puertos y los servicios; y el equipo auditor debe relacionar cada evidencia con un riesgo y una recomendación. De esta manera, el proyecto no solo demuestra vulnerabilidades, sino también cómo deberían gestionarse para reducir el riesgo.",
    )

    polish_document(doc)
    output = OUT / "Establecimiento_Controles_Manual.docx"
    doc.save(output)
    shutil.copyfile(output, OUT / "Establecimiento_Controles.docx")
    return output


if __name__ == "__main__":
    for built in [build_password_policy(), build_secure_use_manual()]:
        print(built)
