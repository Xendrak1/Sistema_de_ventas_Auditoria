from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "documentacion" / "entregables"
EVID = ROOT / "documentacion" / "evidencias"
EVID_DOCX = ROOT / "documentacion" / "evidencias_docx"
OUT.mkdir(parents=True, exist_ok=True)

TEAM = [
    "Rodriguez Soliz Jose Eduardo",
    "Salas Huatta Henry",
    "Sanchez Paniagua Felix Santiago",
    "Taceo Adelia Camila",
    "Tarqui Ignacio Gonzalo",
]

PROJECT_TITLE = "Auditoria de Seguridad y Politicas de Contrasenas para un Sistema de Ventas en un entorno Linux-Fedora 39"

SOURCES = [
    "Flask. (2026). Deploying to Production. https://flask.palletsprojects.com/en/stable/deploying/",
    "Fedora Project. (2026). Fedora Server Documentation. https://docs.fedoraproject.org/en-US/fedora-server/",
    "Lyon, G. F. (2022). Nmap Network Scanning: Official Nmap Project Guide. https://nmap.org/book/man.html",
    "National Institute of Standards and Technology. (2025). Digital Identity Guidelines: Authentication and Lifecycle Management (SP 800-63B). https://csrc.nist.gov/pubs/sp/800/63/b/upd2/final",
    "Open Worldwide Application Security Project. (2021). OWASP Top 10:2021. https://owasp.org/Top10/2021/",
    "Open Worldwide Application Security Project. (2024). Password Storage Cheat Sheet. https://cheatsheetseries.owasp.org/cheatsheets/Password_Storage_Cheat_Sheet.html",
    "SQLite Consortium. (2025). Appropriate Uses For SQLite. https://sqlite.org/whentouse.html",
    "Tenable. (2026). Tenable Nessus Documentation. https://docs.tenable.com/Nessus.htm",
]

RISK_ROWS = [
    ["Sistema web", "Acceso no autorizado", "Credenciales debiles y contrasenas predecibles", "Alto", "Alta", "Critico", "Politica de contrasenas, bloqueo por intentos fallidos y revision de usuarios."],
    ["Sistema web", "Ingreso sin autenticacion valida", "SQL injection en formulario de login", "Alto", "Media", "Alto", "Consultas parametrizadas, validacion de entradas y pruebas de seguridad."],
    ["Sistema web", "Robo de sesion o ejecucion de scripts", "XSS reflejado en busqueda", "Medio", "Media", "Medio", "Escapar salidas HTML, sanitizar entradas y aplicar Content Security Policy."],
    ["Sistema web", "Ejecucion persistente de JavaScript", "XSS almacenado en descripcion de productos", "Alto", "Media", "Alto", "No usar contenido de usuario como seguro, sanitizar y validar campos."],
    ["Base de datos", "Exposicion de credenciales", "Contrasenas almacenadas en texto plano", "Alto", "Media", "Alto", "Hash seguro con bcrypt o argon2 y sal unica por usuario."],
    ["Servidor Fedora", "Enumeracion de servicios", "Puerto 5000 expuesto durante el laboratorio", "Medio", "Alta", "Alto", "Firewall, apertura temporal del puerto y cierre al finalizar la presentacion."],
    ["Servidor Fedora", "Resolucion de nombres innecesaria", "Puerto 5355/LLMNR abierto", "Medio", "Media", "Medio", "Deshabilitar LLMNR si no es requerido para el laboratorio."],
    ["Aplicacion Flask", "Exposicion de informacion tecnica", "Modo debug activo", "Alto", "Media", "Alto", "Desactivar debug fuera del laboratorio y usar servidor WSGI para produccion."],
    ["Aplicacion Flask", "Uso de servidor no productivo", "Werkzeug Development Server expuesto", "Medio", "Alta", "Alto", "Usar Gunicorn/uWSGI detras de Nginx o Apache en entornos reales."],
    ["Servidor Fedora", "Divulgacion de informacion temporal", "ICMP Timestamp Request habilitado", "Bajo", "Media", "Bajo", "Filtrar o deshabilitar respuestas ICMP timestamp si no son necesarias."],
    ["Sistema web", "Debilidad en manejo de sesion", "Cookies expiradas o sin endurecimiento suficiente", "Medio", "Media", "Medio", "Configurar atributos Secure, HttpOnly, SameSite y expiracion adecuada."],
    ["Sistema web", "Acciones no autorizadas", "Formularios sin token CSRF", "Medio", "Media", "Medio", "Implementar CSRF tokens en formularios y validar origen de solicitudes."],
    ["Servidor Fedora", "Uso indebido de privilegios", "Usuarios o servicios con permisos excesivos", "Alto", "Media", "Alto", "Principio de minimo privilegio y cuentas separadas para administracion."],
    ["Sistema web y servidor", "Perdida de disponibilidad", "Ausencia de procedimiento de respaldo", "Alto", "Baja", "Medio", "Copias periodicas de la base SQLite y del codigo fuente versionado."],
]

SPANISH_REPLACEMENTS = {
    "Auditoria": "Auditoría",
    "auditoria": "auditoría",
    "Indice": "Índice",
    "indice": "índice",
    "Gestion": "Gestión",
    "gestion": "gestión",
    "Introduccion": "Introducción",
    "introduccion": "introducción",
    "Fundamentacion": "Fundamentación",
    "fundamentacion": "fundamentación",
    "Descripcion": "Descripción",
    "descripcion": "descripción",
    "Analisis": "Análisis",
    "analisis": "análisis",
    "Politica": "Política",
    "politica": "política",
    "Contrasenas": "Contraseñas",
    "contrasenas": "contraseñas",
    "Codigo": "Código",
    "codigo": "código",
    "academico": "académico",
    "academica": "académica",
    "tecnico": "técnico",
    "tecnica": "técnica",
    "teorica": "teórica",
    "Teorica": "Teórica",
    "Catalogo": "Catálogo",
    "catalogo": "catálogo",
    "aplicacion": "aplicación",
    "Aplicacion": "Aplicación",
    "autenticacion": "autenticación",
    "Autenticacion": "Autenticación",
    "sesion": "sesión",
    "Sesion": "Sesión",
    "informacion": "información",
    "Informacion": "Información",
    "validacion": "validación",
    "Validacion": "Validación",
    "direccion": "dirección",
    "Direccion": "Dirección",
    "Direcciónes": "Direcciones",
    "maquina": "máquina",
    "maquinas": "máquinas",
    "tambien": "también",
    "Tambien": "También",
    "tecnologia": "tecnología",
    "Tecnologia": "Tecnología",
    "deteccion": "detección",
    "Deteccion": "Detección",
    "resolucion": "resolución",
    "Resolucion": "Resolución",
    "recomendacion": "recomendación",
    "Recomendacion": "Recomendación",
    "exposicion": "exposición",
    "Exposicion": "Exposición",
    "produccion": "producción",
    "Produccion": "Producción",
    "publicacion": "publicación",
    "unicamente": "únicamente",
    "despues": "después",
    "titulos": "títulos",
    "subtitulos": "subtítulos",
    "tamano": "tamaño",
    "Tamano": "Tamaño",
    "pagina": "página",
    "Pagina": "Página",
    "modulo": "módulo",
    "Modulo": "Módulo",
    "modulos": "módulos",
    "basico": "básico",
    "basicos": "básicos",
    "debil": "débil",
    "debiles": "débiles",
    "critico": "crítico",
    "criticos": "críticos",
    "parametro": "parámetro",
    "Parametro": "Parámetro",
    "busqueda": "búsqueda",
    "Busqueda": "Búsqueda",
    "podria": "podría",
    "podrian": "podrían",
    "ejecutandose": "ejecutándose",
    "valida": "válida",
    "validas": "válidas",
    "logica": "lógica",
    "especificos": "específicos",
    "disenado": "diseñado",
    "prediseńadas": "prediseñadas",
    "prediseÃ±adas": "prediseñadas",
    "Sistemás": "Sistemas",
    "sistemás": "sistemas",
    "Sistema_de_ventas_Auditoría": "Sistema_de_ventas_Auditoria",
}


def polish_text(text):
    for old, new in SPANISH_REPLACEMENTS.items():
        text = text.replace(old, new)
    return text


def polish_document(doc):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.text = polish_text(run.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.text = polish_text(run.text)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_table(table, header_fill="E8EEF5"):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, row in enumerate(table.rows):
        if i == 0:
            set_repeat_table_header(row)
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(9)
            if i == 0:
                shade_cell(cell, header_fill)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True


def set_doc_style(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 2.0
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Heading 1", 16, "1F4D78"),
        ("Heading 2", 14, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def add_paragraph(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        r1.bold = True
        r2 = p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)


def add_cover(doc, title, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run("Universidad\nMateria: Auditoria de Sistemas\n")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(title)
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(16)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run(subtitle)
    sub_run.font.name = "Times New Roman"
    sub_run.font.size = Pt(12)

    doc.add_paragraph()
    add_paragraph(doc, "Integrantes:")
    for member in TEAM:
        add_bullet(doc, member)
    add_paragraph(doc, "Gestion: 2026")
    add_paragraph(doc, "Repositorio: https://github.com/Xendrak1/Sistema_de_ventas_Auditoria")
    doc.add_page_break()


def add_front_matter(doc):
    doc.add_heading("Indice general", level=1)
    add_paragraph(doc, "Actualizar este indice desde Microsoft Word mediante Referencias > Tabla de contenido, despues de revisar los titulos y subtitulos del documento.")
    doc.add_heading("Indice de tablas", level=1)
    tables = [
        "Tabla 1. Rutas principales del sistema web.",
        "Tabla 2. Tablas principales de la base de datos SQLite.",
        "Tabla 3. Herramientas utilizadas en el proyecto.",
        "Tabla 4. Resumen de evidencias integradas al informe.",
        "Tabla 5. Interpretación de puertos y servicios detectados con Nmap.",
        "Tabla 6. Interpretación de hallazgos detectados por Nessus.",
        "Tabla 7. Relación entre fallas prediseñadas, evidencia y control propuesto.",
        "Tabla 8. Activos evaluados.",
        "Tabla 9. Amenazas del sistema web y servidor.",
        "Tabla 10. Vulnerabilidades identificadas.",
        "Tabla 11. Matriz de análisis de riesgos.",
        "Tabla 12. Controles recomendados.",
    ]
    for item in tables:
        add_bullet(doc, item)
    doc.add_heading("Indice de figuras", level=1)
    figures = [
        "Figura 1. Pantalla de inicio de sesion.",
        "Figura 2. Panel principal del sistema de ventas.",
        "Figura 3. Catalogo de productos.",
        "Figura 4. Modulo de reportes.",
        "Figura 5. Direcciones IP del laboratorio.",
        "Figura 6. Servicio Flask ejecutandose en Fedora.",
        "Figura 7. Acceso al login desde Kali Linux.",
        "Figura 8. Panel principal accesible desde Kali.",
        "Figura 9. Modulo de clientes.",
        "Figura 10. Busqueda de productos.",
        "Figura 11. Modulo de reportes.",
        "Figura 12. Resultado de Nmap con deteccion de servicio.",
        "Figura 13. Escaneo de puertos desde Kali.",
        "Figura 14. Confirmacion de servicios desde Fedora.",
        "Figura 15. Resultado de Nessus sin autenticacion.",
        "Figura 16. Prueba de SQL Injection en login.",
        "Figura 17. Prueba de XSS reflejado.",
        "Figura 18. Prueba de XSS almacenado.",
        "Figura 19. Contrasenas almacenadas en texto plano.",
    ]
    for item in figures:
        add_bullet(doc, item)
    doc.add_page_break()


def add_matrix_table(doc, rows):
    headers = ["Activo", "Amenaza", "Vulnerabilidad", "Impacto", "Prob.", "Riesgo", "Control recomendado"]
    table = doc.add_table(rows=1, cols=len(headers))
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    style_table(table)
    add_caption(doc, "Tabla 11. Matriz de analisis de riesgos.")


def add_tool_table(doc):
    data = [
        ["VirtualBox", "Creacion de maquina virtual Fedora 39 para laboratorio controlado."],
        ["Fedora 39", "Sistema operativo del servidor auditado."],
        ["Python y Flask", "Desarrollo y ejecucion del sistema web de ventas."],
        ["SQLite", "Base de datos local simple para usuarios, productos, clientes y ventas."],
        ["Git y GitHub", "Control de versiones, respaldo de codigo y evidencia de entrega."],
        ["Nmap", "Escaneo de puertos, servicios y version del servidor."],
        ["Nessus Essentials", "Analisis de vulnerabilidades y generacion de reporte tecnico."],
        ["Microsoft Word", "Redaccion del informe con formato APA, tablas, figuras y referencias."],
    ]
    table = doc.add_table(rows=1, cols=2)
    set_cell_text(table.rows[0].cells[0], "Herramienta", bold=True)
    set_cell_text(table.rows[0].cells[1], "Uso en el proyecto", bold=True)
    for item in data:
        cells = table.add_row().cells
        set_cell_text(cells[0], item[0])
        set_cell_text(cells[1], item[1])
    style_table(table)
    add_caption(doc, "Tabla 3. Herramientas utilizadas en el proyecto.")


def add_evidence_summary_table(doc):
    data = [
        ["Direcciones IP", "Fedora 192.168.56.103 y Kali 192.168.56.102", "Demuestra que ambas maquinas estan en la misma red de laboratorio."],
        ["Servicio Flask", "python app.py ejecutandose en Fedora", "Demuestra que el sistema fue desplegado en el servidor Linux."],
        ["Acceso desde Kali", "Login y panel principal desde navegador de Kali", "Demuestra que el sistema es accesible desde una maquina auditora externa."],
        ["Nmap", "Puerto 5000/TCP abierto con Werkzeug/Python", "Demuestra la exposicion del servicio web y permite identificar tecnologia usada."],
        ["Nessus", "Escaneo sin autenticacion contra Fedora", "Demuestra analisis automatizado de infraestructura y servicio HTTP."],
        ["Pruebas manuales", "SQL Injection, XSS y contrasenas en texto plano", "Demuestra fallas prediseñadas propias de la aplicacion web."],
    ]
    table = doc.add_table(rows=1, cols=3)
    for idx, header in enumerate(["Evidencia", "Resultado observado", "Interpretacion"]):
        set_cell_text(table.rows[0].cells[idx], header, True)
    for row in data:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    style_table(table)
    add_caption(doc, "Tabla 4. Resumen de evidencias integradas al informe.")


def add_nmap_interpretation_table(doc):
    data = [
        ["5000/TCP", "Abierto", "HTTP Werkzeug/Python", "Es el sistema web de ventas publicado en Flask.", "Un atacante sabría dónde está la aplicación y qué tecnología usa para buscar fallas específicas.", "Mantener el puerto abierto solo durante la práctica y usar un servidor WSGI en un entorno real."],
        ["5355/TCP", "Abierto", "LLMNR / systemd-resolved", "Servicio de resolución de nombres local del sistema operativo.", "Puede facilitar reconocimiento interno o ataques de suplantación en redes mal configuradas.", "Deshabilitar LLMNR si no es necesario para el laboratorio."],
        ["/login", "Ruta detectada", "HTTP", "El sistema redirige al formulario de autenticación.", "Permite ubicar el punto principal de ataque contra credenciales y autenticación.", "Endurecer el login con bloqueo por intentos, validaciones y registro de eventos."],
        ["Ingreso - Sistema de Venta", "Título web", "Aplicación Flask", "Nmap identificó el título de la página.", "Revela información útil para perfilar el sistema durante la fase de reconocimiento.", "Evitar información innecesaria sobre tecnología o entorno."],
    ]
    table = doc.add_table(rows=1, cols=6)
    for idx, header in enumerate(["Elemento", "Estado", "Servicio", "Qué significa", "Cómo podría aprovecharse", "Control recomendado"]):
        set_cell_text(table.rows[0].cells[idx], header, True)
    for row in data:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    style_table(table)
    add_caption(doc, "Tabla 5. Interpretación de puertos y servicios detectados con Nmap.")


def add_nessus_findings_table(doc):
    data = [
        ["HTTP (Multiple Issues)", "Informativo", "Nessus agrupó observaciones del servicio HTTP expuesto.", "Puede orientar a un atacante sobre configuraciones web débiles o puntos que requieren revisión.", "Revisar cabeceras HTTP, cookies, métodos permitidos y configuración del servidor."],
        ["Python Remote HTTP Detection", "Informativo", "Detectó que el servicio web está relacionado con Python.", "La tecnología detectada permite buscar vulnerabilidades o malas prácticas propias del stack Python/Flask.", "Ocultar información innecesaria y usar configuración de despliegue adecuada."],
        ["Web Application Cookies Are Expired", "Informativo", "Se observaron cookies con expiración o manejo débil.", "Un manejo incorrecto de cookies puede afectar sesiones y autenticación si se combina con otras fallas.", "Configurar expiración, HttpOnly, SameSite y Secure cuando corresponda."],
        ["ICMP Timestamp Request Remote Date Disclosure", "Bajo", "El servidor responde solicitudes ICMP timestamp.", "Puede ayudar a conocer información de tiempo del host y apoyar tareas de reconocimiento.", "Filtrar o deshabilitar respuestas ICMP timestamp si no son necesarias."],
        ["Service / OS Detection", "Informativo", "Nessus identificó servicios y huellas del sistema.", "Ayuda a perfilar el servidor antes de intentar ataques más específicos.", "Reducir servicios expuestos y mantener el sistema actualizado."],
    ]
    table = doc.add_table(rows=1, cols=5)
    for idx, header in enumerate(["Hallazgo de Nessus", "Nivel", "Qué indica", "Cómo podría aprovecharse", "Control recomendado"]):
        set_cell_text(table.rows[0].cells[idx], header, True)
    for row in data:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    style_table(table)
    add_caption(doc, "Tabla 6. Interpretación de hallazgos detectados por Nessus.")


def add_manual_findings_table(doc):
    data = [
        ["Credenciales débiles", "Ingreso con admin/admin123", "Alto", "Aplicar política de contraseñas y bloqueo por intentos fallidos."],
        ["SQL Injection", "Ingreso con payload ' OR '1'='1' --", "Alto", "Usar consultas parametrizadas y validación de entradas."],
        ["XSS reflejado", "Parámetro q de búsqueda ejecuta JavaScript", "Medio", "Escapar salidas HTML y aplicar Content Security Policy."],
        ["XSS almacenado", "Descripción de producto almacena script", "Alto", "Sanitizar contenido de usuario antes de guardar y mostrar."],
        ["Contraseñas en texto plano", "Consulta SQLite muestra admin123 y ventas123", "Alto", "Guardar contraseñas con hash seguro como bcrypt o argon2."],
        ["Servidor de desarrollo", "Werkzeug expuesto en puerto 5000", "Medio", "Usar Gunicorn/uWSGI con Nginx o Apache fuera del laboratorio."],
    ]
    table = doc.add_table(rows=1, cols=4)
    for idx, header in enumerate(["Falla prediseñada", "Evidencia", "Nivel", "Control propuesto"]):
        set_cell_text(table.rows[0].cells[idx], header, True)
    for row in data:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    style_table(table)
    add_caption(doc, "Tabla 7. Relación entre fallas prediseñadas, evidencia y control propuesto.")


def add_image(doc, filename, caption):
    path = EVID / filename
    if path.exists():
        doc.add_picture(str(path), width=Inches(5.8))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_caption(doc, caption)


def add_docx_evidence_image(doc, filename, caption, description=None, width=5.9):
    path = EVID_DOCX / filename
    if path.exists():
        if description:
            add_paragraph(doc, description)
        doc.add_picture(str(path), width=Inches(width))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_caption(doc, caption)


def build_main_report():
    doc = Document()
    set_doc_style(doc)
    add_cover(doc, PROJECT_TITLE, "Informe principal de auditoria")
    add_front_matter(doc)

    doc.add_heading("Introduccion", level=1)
    add_paragraph(doc, "Este informe presenta la auditoría de seguridad realizada a un sistema web de ventas desarrollado en Python con Flask y desplegado sobre Fedora 39. El proyecto fue construido como un laboratorio universitario: el sistema funciona, tiene código fuente propio y además incluye fallas prediseñadas para poder analizarlas, explicarlas y proponer controles de mejora.")
    add_paragraph(doc, "El escenario se trabajó con dos máquinas virtuales. Fedora 39 actúa como servidor, porque ahí se ejecuta la aplicación y la base de datos SQLite. Kali Linux actúa como equipo auditor, desde donde se accede a la página, se ejecuta Nmap y se realizan pruebas manuales. Nessus complementa el trabajo con un escaneo automatizado del servidor y del servicio web expuesto.")
    add_paragraph(doc, "La idea no fue publicar una aplicación vulnerable en internet, sino crear un ambiente controlado para aprender el proceso completo de una auditoría: desplegar un servicio, reconocer puertos, revisar hallazgos, comprobar vulnerabilidades, valorar riesgos y plantear controles. Por eso las pruebas se realizaron dentro de una red de laboratorio en VirtualBox.")
    add_paragraph(doc, "El sistema incluye módulos de login, productos, clientes, ventas, reportes y búsqueda. Sus fallas principales son credenciales débiles, contraseñas en texto plano, SQL Injection, XSS reflejado, XSS almacenado, falta de CSRF y uso del servidor de desarrollo de Flask. Estas vulnerabilidades permiten relacionar la teoría de seguridad web con evidencias reales obtenidas durante la práctica.")
    add_paragraph(doc, "Además del análisis técnico, el trabajo incluye una matriz de riesgos, recomendaciones, controles de seguridad y una política de contraseñas. De esta forma se cubre lo solicitado para la materia: código fuente, despliegue, auditoría del sistema, evidencias, análisis de riesgo y documentación de controles.")

    doc.add_heading("Fundamentacion teorica", level=1)
    add_paragraph(doc, "La seguridad de aplicaciones web busca proteger la confidencialidad, integridad y disponibilidad de los datos procesados por un sistema. OWASP identifica categorias de riesgo como control de acceso roto, fallas criptograficas e inyeccion, las cuales son relevantes para sistemas que gestionan usuarios, autenticacion y operaciones comerciales.")
    add_paragraph(doc, "Las politicas de contrasenas son un componente esencial de la autenticacion. NIST SP 800-63B recomienda controles orientados a longitud adecuada, proteccion contra secretos comprometidos y limitacion de intentos fallidos. En el proyecto, estas recomendaciones se traducen en controles propuestos para el sistema web y el servidor.")
    add_paragraph(doc, "Nmap es una herramienta de exploracion de red que permite identificar hosts, puertos abiertos, servicios y versiones. Nessus Essentials, por su parte, permite realizar evaluaciones de vulnerabilidades sobre hosts y servicios, generando hallazgos con severidad y recomendaciones.")
    add_paragraph(doc, "SQLite fue seleccionado como base de datos local por su simplicidad y porque no requiere un servidor de base de datos independiente. Esta decision es adecuada para un laboratorio academico y para un sistema pequeño, siempre que no se presente como arquitectura final de produccion.")

    doc.add_heading("Objetivos", level=1)
    doc.add_heading("Objetivo general", level=2)
    add_paragraph(doc, "Realizar una auditoria de seguridad a un sistema web de ventas desplegado sobre Fedora 39, identificando vulnerabilidades, amenazas y riesgos, y proponiendo controles de seguridad enfocados en autenticacion, contrasenas y proteccion del servidor.")
    doc.add_heading("Objetivos especificos", level=2)
    objectives = [
        "Desarrollar un sistema web de ventas simple que permita operaciones basicas de productos, clientes, ventas y reportes.",
        "Desplegar el sistema en un entorno Fedora 39 sobre VirtualBox para simular un servidor Linux.",
        "Ejecutar analisis de red y servicios con Nmap.",
        "Ejecutar analisis de vulnerabilidades con Nessus Essentials.",
        "Identificar amenazas y vulnerabilidades asociadas al sistema web y al servidor.",
        "Elaborar una matriz de riesgos que relacione activos, amenazas, vulnerabilidades, impacto, probabilidad y controles.",
        "Proponer politicas de contrasenas y controles de seguridad aplicables al sistema y servidor.",
    ]
    for item in objectives:
        add_bullet(doc, item)

    doc.add_heading("Alcance del proyecto", level=1)
    add_paragraph(doc, "El alcance incluye el codigo fuente del sistema de ventas, la base de datos SQLite local, el servidor Fedora 39 utilizado en la demostracion, el puerto de aplicacion Flask, la autenticacion del sistema y los controles administrativos propuestos. No incluye explotacion ofensiva fuera del laboratorio, publicacion permanente en internet ni tratamiento de datos reales de clientes.")

    doc.add_heading("Normas APA aplicadas", level=1)
    add_paragraph(doc, "El documento fue preparado con criterios basicos de presentacion academica: hoja tamano carta, margenes de 2.54 cm, tipo de letra Times New Roman, tamano 12 para el cuerpo del texto, titulos jerarquicos, tablas numeradas, figuras numeradas y bibliografia al final del documento.")
    add_paragraph(doc, "Para completar el formato en Microsoft Word, se recomienda actualizar el indice general desde Referencias > Tabla de contenido, generar el indice de tablas y figuras desde Referencias > Insertar tabla de ilustraciones, y registrar las fuentes desde Referencias > Insertar cita si el docente exige bibliografia automatica.")

    doc.add_heading("Descripcion del sistema de ventas", level=1)
    add_paragraph(doc, "El sistema de ventas permite iniciar sesion, visualizar un panel principal, registrar productos, registrar clientes, registrar ventas, consultar reportes y buscar productos. La interfaz fue fortalecida con tarjetas de productos, imagenes locales, SKU, categorias, alertas de inventario y una pantalla de reportes.")
    add_paragraph(doc, "La aplicacion utiliza Flask como framework web, SQLite como base de datos local y archivos estaticos para estilos e imagenes. El repositorio oficial del proyecto se encuentra en GitHub para facilitar el despliegue y la trazabilidad de cambios.")

    doc.add_heading("Desarrollo del servicio web de ventas", level=1)
    add_paragraph(doc, "El servicio web fue desarrollado como una aplicacion monolitica simple en Flask. El archivo principal app.py define la configuracion de la aplicacion, la conexion a SQLite, la inicializacion de tablas, las rutas web y la carga de plantillas HTML.")
    route_table = doc.add_table(rows=1, cols=3)
    for idx, header in enumerate(["Ruta", "Modulo", "Descripcion"]):
        set_cell_text(route_table.rows[0].cells[idx], header, True)
    for row in [
        ["/login", "Autenticacion", "Permite el ingreso de usuarios al sistema."],
        ["/dashboard", "Panel", "Muestra metricas generales, productos destacados y alertas de stock."],
        ["/productos", "Inventario", "Lista productos con imagen, SKU, categoria, precio y stock."],
        ["/clientes", "Clientes", "Permite registrar y listar clientes."],
        ["/ventas", "Ventas", "Permite registrar operaciones de venta y calcular totales."],
        ["/reportes", "Reportes", "Presenta ventas por producto y control de inventario."],
        ["/buscar", "Busqueda", "Permite consultar productos por nombre o descripcion."],
    ]:
        cells = route_table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    style_table(route_table)
    add_caption(doc, "Tabla 1. Rutas principales del sistema web.")

    db_table = doc.add_table(rows=1, cols=3)
    for idx, header in enumerate(["Tabla", "Campos principales", "Uso"]):
        set_cell_text(db_table.rows[0].cells[idx], header, True)
    for row in [
        ["usuarios", "usuario, password, rol", "Control de acceso al sistema."],
        ["productos", "sku, categoria, nombre, descripcion, precio, stock, imagen", "Catalogo e inventario."],
        ["clientes", "nombre, telefono, correo", "Datos basicos de clientes."],
        ["ventas", "cliente_id, producto_id, cantidad, total, fecha", "Registro de transacciones."],
    ]:
        cells = db_table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    style_table(db_table)
    add_caption(doc, "Tabla 2. Tablas principales de la base de datos SQLite.")

    add_paragraph(doc, "Desde el punto de vista de auditoria, el sistema incluye vulnerabilidades intencionales: credenciales debiles, contrasenas en texto plano, SQL injection en login y busqueda, XSS reflejado, XSS almacenado, ausencia de CSRF, clave de sesion debil y modo debug activo. Estas fallas permiten justificar tecnicamente los controles propuestos.")
    add_image(doc, "figura-01-login.png", "Figura 1. Pantalla de inicio de sesion.")
    add_image(doc, "figura-02-dashboard.png", "Figura 2. Panel principal del sistema de ventas.")
    add_image(doc, "figura-03-productos.png", "Figura 3. Catalogo de productos.")
    add_image(doc, "figura-04-reportes.png", "Figura 4. Modulo de reportes.")

    doc.add_heading("Herramientas utilizadas", level=1)
    add_tool_table(doc)

    doc.add_heading("Despliegue del servidor Linux Fedora 39", level=1)
    add_paragraph(doc, "El despliegue recomendado se realiza en VirtualBox con Fedora 39 y adaptador de red en modo puente. Este modo permite que el servidor obtenga una direccion IP visible dentro de la red local, facilitando las pruebas desde Nmap y Nessus.")
    for step in [
        "Crear una maquina virtual en VirtualBox e instalar Fedora 39.",
        "Actualizar el sistema con sudo dnf update -y.",
        "Instalar dependencias con sudo dnf install -y git python3 python3-pip.",
        "Clonar el repositorio desde GitHub.",
        "Crear un entorno virtual de Python.",
        "Instalar dependencias desde requirements.txt.",
        "Ejecutar python app.py para iniciar el servicio en el puerto 5000.",
        "Abrir temporalmente el puerto 5000/tcp en firewalld si se requiere acceso desde otra maquina.",
    ]:
        add_number(doc, step)

    doc.add_heading("Comandos de despliegue", level=2)
    commands = [
        "sudo dnf update -y",
        "sudo dnf install -y git python3 python3-pip",
        "git clone https://github.com/Xendrak1/Sistema_de_ventas_Auditoria.git",
        "cd Sistema_de_ventas_Auditoria",
        "python3 -m venv .venv",
        "source .venv/bin/activate",
        "pip install -r requirements.txt",
        "python app.py",
    ]
    for command in commands:
        add_paragraph(doc, command)

    doc.add_heading("Despliegue del servicio web", level=1)
    add_paragraph(doc, "El servicio web se ejecuta en http://IP_DEL_SERVIDOR:5000. Para la presentacion, el servidor debe permanecer activo unicamente durante la demostracion. Al finalizar, se recomienda detener Flask con Ctrl + C o apagar la maquina virtual.")
    add_paragraph(doc, "Credenciales de prueba: usuario admin y contrasena admin123. Estas credenciales son intencionalmente debiles para demostrar la necesidad de politicas de contrasenas.")

    doc.add_heading("Evidencias del despliegue y auditoria", level=1)
    add_paragraph(doc, "Las siguientes evidencias fueron obtenidas durante la ejecucion del laboratorio. Su objetivo es demostrar que el sistema no solo existe como codigo fuente, sino que fue desplegado en Fedora 39, accedido desde Kali Linux y analizado con herramientas de auditoria. La secuencia de capturas permite reconstruir el proceso completo: configuracion de red, ejecucion del servicio, acceso al sistema, reconocimiento con Nmap, validacion con Nessus y pruebas manuales de fallas prediseñadas.")
    add_evidence_summary_table(doc)

    doc.add_heading("Evidencias de red y despliegue", level=2)
    add_docx_evidence_image(
        doc,
        "evidencia_docx_03.png",
        "Figura 5. Direcciones IP del laboratorio.",
        "La evidencia muestra las interfaces de red configuradas en el laboratorio. Fedora utiliza la direccion 192.168.56.103 y Kali Linux la direccion 192.168.56.102 dentro de la red solo anfitrion. Esto confirma que ambas maquinas se encuentran en el mismo segmento de red y que Kali puede actuar como equipo auditor contra el servidor Fedora.",
    )
    add_docx_evidence_image(
        doc,
        "evidencia_docx_11.png",
        "Figura 6. Servicio Flask ejecutandose en Fedora.",
        "La captura evidencia la ejecucion de python app.py en Fedora. El servidor Flask queda escuchando en el puerto 5000 y registra solicitudes HTTP provenientes de la maquina auditora. Esta prueba demuestra que el sistema fue desplegado correctamente en el servidor Linux y que no se trata de una pagina estatica o de una simulacion.",
    )
    add_docx_evidence_image(
        doc,
        "evidencia_docx_05.png",
        "Figura 7. Acceso al login desde Kali Linux.",
        "La pantalla de login abierta desde Kali Linux demuestra que el servicio web alojado en Fedora es accesible desde una maquina externa. Esta evidencia es importante porque confirma la exposicion real del servicio dentro de la red de laboratorio y justifica el uso de herramientas de reconocimiento como Nmap y Nessus.",
    )

    doc.add_heading("Evidencias funcionales del sistema", level=2)
    add_docx_evidence_image(
        doc,
        "evidencia_docx_01.png",
        "Figura 8. Panel principal accesible desde Kali.",
        "El panel principal muestra metricas del sistema de ventas, productos destacados y alertas de inventario. Esta captura prueba que el inicio de sesion fue exitoso y que Kali no solo accede a la pantalla inicial, sino tambien al contenido interno del sistema autenticado.",
    )
    add_docx_evidence_image(
        doc,
        "evidencia_docx_06.png",
        "Figura 9. Modulo de clientes.",
        "El modulo de clientes permite registrar y listar informacion basica. En el contexto de auditoria, este modulo representa un activo de informacion porque procesa datos de clientes y requiere controles de acceso, validacion de entradas y proteccion contra acciones no autorizadas.",
    )
    add_docx_evidence_image(
        doc,
        "evidencia_docx_02.png",
        "Figura 10. Busqueda de productos.",
        "La busqueda de productos permite consultar registros del catalogo. Este punto es relevante porque la entrada del usuario se utiliza en consultas y en la respuesta HTML, por lo que se convierte en una zona de prueba para SQL Injection y XSS reflejado.",
    )
    add_docx_evidence_image(
        doc,
        "evidencia_docx_12.png",
        "Figura 11. Modulo de reportes.",
        "El modulo de reportes presenta ventas por producto y control de stock. Esta evidencia muestra que el sistema cuenta con funcionalidad administrativa suficiente para justificar controles de acceso, integridad de datos y respaldo de informacion.",
    )

    doc.add_heading("Análisis con Nmap", level=1)
    add_paragraph(doc, "El análisis con Nmap se ejecutó desde Kali Linux hacia la IP del servidor Fedora. La finalidad fue confirmar qué servicios estaban abiertos y comprobar que el sistema de ventas realmente estaba publicado en la red de laboratorio.")
    for command in ["nmap -sV -p 5000 192.168.56.103", "nmap -A 192.168.56.103", "nmap -p- 192.168.56.103"]:
        add_paragraph(doc, command)
    add_paragraph(doc, "El hallazgo más importante fue el puerto 5000/TCP abierto con un servicio HTTP Werkzeug/Python. Eso corresponde a la aplicación Flask del sistema de ventas. También apareció el puerto 5355/TCP, relacionado con LLMNR, un servicio de resolución de nombres local. Nmap no explota vulnerabilidades por sí solo en este caso, pero ayuda a descubrir qué está expuesto y por dónde podría empezar un atacante.")
    add_docx_evidence_image(
        doc,
        "evidencia_docx_07.png",
        "Figura 12. Resultado de Nmap con deteccion de servicio.",
        "La captura muestra el resultado de un escaneo Nmap con deteccion de version. Se observa el puerto 5000/TCP abierto, el servicio HTTP Werkzeug httpd 3.1.8 y Python 3.12.0. Tambien se identifica la ruta /login, lo que confirma que el puerto corresponde al sistema web de ventas.",
    )
    add_docx_evidence_image(
        doc,
        "evidencia_docx_09.png",
        "Figura 13. Escaneo de puertos desde Kali.",
        "Este escaneo muestra puertos abiertos adicionales, entre ellos 5000/TCP y 5355/TCP. El puerto 5000 corresponde al sistema web Flask, mientras que el puerto 5355 se relaciona con LLMNR, un servicio de resolucion de nombres local del sistema operativo.",
    )
    add_docx_evidence_image(
        doc,
        "evidencia_docx_14.png",
        "Figura 14. Confirmacion de servicios desde Fedora.",
        "Desde Fedora se ejecuto una validacion local de puertos y procesos. La evidencia confirma que el puerto 5000 esta asociado al proceso python que ejecuta el sistema web, mientras que el puerto 5355 corresponde a systemd-resolved. Esta comparacion fortalece el resultado de Nmap porque relaciona el hallazgo externo con el proceso real del servidor.",
    )
    add_nmap_interpretation_table(doc)

    doc.add_heading("Análisis con Nessus", level=1)
    add_paragraph(doc, "El escaneo con Nessus se ejecutó sin autenticación contra el servidor Fedora. Esto significa que Nessus revisó lo que cualquier equipo de la red podía ver desde afuera: servicios abiertos, respuestas HTTP, detección de tecnologías y configuraciones generales. Las fallas internas de la aplicación, como SQL Injection o XSS dentro del sistema, se comprobaron aparte mediante pruebas manuales.")
    add_docx_evidence_image(
        doc,
        "evidencia_docx_18.png",
        "Figura 15. Resultado de Nessus sin autenticación.",
        "La captura muestra hallazgos informativos y de baja severidad generados por Nessus, como detección HTTP, detección remota de Python, información del sistema operativo, cookies expiradas y divulgación de timestamp ICMP. Aunque no todos los hallazgos son críticos, sirven para entender qué información entrega el servidor cuando alguien lo analiza desde la red.",
    )
    add_nessus_findings_table(doc)

    doc.add_heading("Pruebas manuales de vulnerabilidades prediseñadas", level=1)
    add_paragraph(doc, "Las pruebas manuales se realizaron para validar fallas propias de la lógica de la aplicación. Este paso es necesario porque herramientas como Nessus no siempre detectan vulnerabilidades que están detrás del login o que dependen de formularios específicos. Aquí se comprobaron las fallas que fueron diseñadas intencionalmente para el proyecto.")
    add_docx_evidence_image(
        doc,
        "evidencia_docx_10.png",
        "Figura 16. Prueba de SQL Injection en login.",
        "La captura muestra el uso del payload ' OR '1'='1' -- en el formulario de inicio de sesión. La prueba evidencia que el login concatena entradas del usuario dentro de una consulta SQL. Si esta falla existiera en un sistema real, una persona podría ingresar sin conocer una contraseña válida.",
    )
    add_docx_evidence_image(
        doc,
        "evidencia_docx_13.png",
        "Figura 17. Prueba de XSS reflejado.",
        "La alerta en el navegador evidencia que la búsqueda refleja entrada del usuario sin escape seguro. En un caso real, esta falla podría usarse para ejecutar JavaScript en el navegador de un usuario si se le envía una URL manipulada.",
    )
    add_docx_evidence_image(
        doc,
        "evidencia_docx_16.png",
        "Figura 18. Registro de payload para XSS almacenado.",
        "En esta captura se observa el registro de un producto cuya descripción contiene un script. La prueba representa XSS almacenado porque el contenido malicioso se guarda en la base de datos y luego vuelve a mostrarse en la interfaz.",
    )
    add_docx_evidence_image(
        doc,
        "evidencia_docx_21.png",
        "Figura 19. Ejecución de XSS almacenado.",
        "La alerta confirma que el script almacenado en la descripción del producto se ejecuta al visualizar el módulo correspondiente. Este hallazgo tiene mayor impacto que un XSS reflejado porque permanece en el sistema hasta que el registro sea corregido o eliminado.",
    )
    add_docx_evidence_image(
        doc,
        "evidencia_docx_08.png",
        "Figura 20. Contraseñas almacenadas en texto plano.",
        "La consulta realizada sobre SQLite muestra usuarios y contraseñas directamente legibles, como admin123 y ventas123. Esta evidencia demuestra una falla crítica de almacenamiento de credenciales: si alguien obtiene la base de datos, también obtiene las contraseñas.",
    )
    add_manual_findings_table(doc)

    doc.add_heading("Activos evaluados", level=1)
    table = doc.add_table(rows=1, cols=3)
    for idx, header in enumerate(["Activo", "Descripcion", "Importancia"]):
        set_cell_text(table.rows[0].cells[idx], header, True)
    for row in [
        ["Sistema web", "Aplicacion Flask de ventas.", "Gestiona acceso, productos, clientes y ventas."],
        ["Base SQLite", "Archivo ventas.db generado por la aplicacion.", "Almacena usuarios y datos operativos."],
        ["Servidor Fedora", "VM Linux usada para despliegue.", "Aloja el servicio web y expone el puerto 5000."],
        ["Repositorio GitHub", "Codigo fuente y documentacion.", "Evidencia de versionamiento y entrega."],
    ]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    style_table(table)
    add_caption(doc, "Tabla 8. Activos evaluados.")

    doc.add_heading("Amenazas del sistema web y servidor", level=1)
    table = doc.add_table(rows=1, cols=3)
    for idx, header in enumerate(["Amenaza", "Descripcion", "Activo afectado"]):
        set_cell_text(table.rows[0].cells[idx], header, True)
    for row in [
        ["Acceso no autorizado", "Ingreso al sistema mediante credenciales debiles o bypass.", "Sistema web"],
        ["Inyeccion SQL", "Manipulacion de consultas mediante entradas maliciosas.", "Sistema web y base de datos"],
        ["XSS", "Ejecucion de scripts en el navegador del usuario.", "Sistema web"],
        ["Enumeracion de servicios", "Identificacion de puertos y servicios expuestos.", "Servidor Fedora"],
        ["Exposicion de informacion", "Divulgacion de errores, debug o datos sensibles.", "Aplicacion Flask"],
    ]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    style_table(table)
    add_caption(doc, "Tabla 9. Amenazas del sistema web y servidor.")

    doc.add_heading("Vulnerabilidades del sistema web y servidor", level=1)
    table = doc.add_table(rows=1, cols=4)
    for idx, header in enumerate(["Vulnerabilidad", "Ubicacion", "Impacto", "Evidencia esperada"]):
        set_cell_text(table.rows[0].cells[idx], header, True)
    for row in [
        ["Credenciales debiles", "Login", "Acceso no autorizado", "Ingreso con admin/admin123."],
        ["Contrasenas en texto plano", "ventas.db", "Exposicion de credenciales", "Revision de la tabla usuarios."],
        ["SQL injection", "/login y /buscar", "Bypass o exposicion de datos", "Prueba con payload controlado."],
        ["XSS reflejado", "/buscar?q=", "Ejecucion de scripts", "Parametro q sin escape seguro."],
        ["XSS almacenado", "Descripcion de productos", "Persistencia de scripts", "Uso de descripcion con HTML no sanitizado."],
        ["Debug activo", "app.py", "Exposicion tecnica", "Ejecucion Flask con debug=True."],
    ]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    style_table(table)
    add_caption(doc, "Tabla 10. Vulnerabilidades identificadas.")

    doc.add_heading("Tabla de analisis de riesgos", level=1)
    add_paragraph(doc, "La matriz de riesgos relaciona activos, amenazas, vulnerabilidades, impacto, probabilidad y controles. Para priorizar, se clasifica el riesgo como Critico, Alto, Medio o Bajo.")
    add_matrix_table(doc, RISK_ROWS)

    doc.add_heading("Controles recomendados", level=1)
    table = doc.add_table(rows=1, cols=4)
    for idx, header in enumerate(["Control", "Tipo", "Descripcion", "Responsable"]):
        set_cell_text(table.rows[0].cells[idx], header, True)
    for row in [
        ["Politica de contrasenas", "Preventivo", "Define longitud, complejidad, bloqueo y recuperacion segura.", "Administrador del sistema"],
        ["Hash de contrasenas", "Preventivo", "Uso de bcrypt o argon2 en lugar de texto plano.", "Desarrollador"],
        ["Consultas parametrizadas", "Preventivo", "Evita SQL injection en login y busqueda.", "Desarrollador"],
        ["Sanitizacion de entradas", "Preventivo", "Reduce XSS reflejado y almacenado.", "Desarrollador"],
        ["Firewall", "Preventivo", "Apertura temporal solo del puerto requerido.", "Administrador del servidor"],
        ["Registro de eventos", "Detectivo", "Logs de login, errores y acciones relevantes.", "Administrador"],
        ["Backups", "Correctivo", "Respaldo de ventas.db y repositorio.", "Equipo del proyecto"],
    ]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    style_table(table)
    add_caption(doc, "Tabla 12. Controles recomendados.")

    doc.add_heading("Conclusiones", level=1)
    conclusions = [
        "El proyecto permite demostrar el ciclo basico de auditoria de seguridad: despliegue, reconocimiento, analisis, identificacion de riesgos y propuesta de controles.",
        "El uso de Fedora 39 en VirtualBox es adecuado para mantener el alcance controlado y evitar la exposicion permanente de vulnerabilidades en internet.",
        "Las vulnerabilidades intencionales permiten evidenciar riesgos comunes en sistemas web, especialmente autenticacion debil, SQL injection, XSS y almacenamiento inseguro de contrasenas.",
        "Nmap y Nessus complementan el analisis al identificar servicios expuestos, configuraciones inseguras y posibles vulnerabilidades del entorno.",
        "La aplicacion requiere controles correctivos antes de cualquier uso real, principalmente hash de contrasenas, consultas parametrizadas, proteccion CSRF, desactivacion de debug y configuracion segura del servidor.",
    ]
    for item in conclusions:
        add_bullet(doc, item)

    doc.add_heading("Bibliografia", level=1)
    for source in SOURCES:
        add_paragraph(doc, source)

    polish_document(doc)
    path = OUT / "Informe_Principal_Auditoria.docx"
    doc.save(path)
    return path


def build_password_policy():
    doc = Document()
    set_doc_style(doc)
    add_cover(doc, "Politica de Uso de Contrasenas", PROJECT_TITLE)

    doc.add_heading("Objetivo", level=1)
    add_paragraph(doc, "Establecer reglas y controles para la creacion, uso, almacenamiento, recuperacion y administracion de contrasenas en el sistema web de ventas y en el servidor Fedora 39.")

    doc.add_heading("Alcance", level=1)
    add_paragraph(doc, "Esta politica aplica a usuarios del sistema de ventas, administradores del servidor, desarrolladores y cualquier integrante que tenga acceso al repositorio, la aplicacion o la maquina virtual utilizada durante la auditoria.")

    doc.add_heading("Responsables", level=1)
    for item in [
        "Administrador del sistema: crear usuarios, revocar accesos y controlar privilegios.",
        "Desarrollador: implementar validaciones, hash de contrasenas y controles de autenticacion.",
        "Usuarios: proteger sus credenciales y reportar accesos sospechosos.",
        "Equipo auditor: verificar el cumplimiento de la politica durante las pruebas.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("Reglas para contrasenas", level=1)
    for item in [
        "La longitud minima recomendada para usuarios del sistema sera de 10 caracteres.",
        "La contrasena debe combinar letras mayusculas, minusculas, numeros y caracteres especiales.",
        "No se permite usar contrasenas comunes como admin123, ventas123, password, 123456 o datos personales.",
        "Las contrasenas no deben compartirse por mensajeria, correo o documentos sin proteccion.",
        "El sistema debe bloquear temporalmente la cuenta despues de 5 intentos fallidos.",
        "La recuperacion de contrasena debe realizarse mediante un procedimiento controlado y registrado.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("Almacenamiento seguro", level=1)
    add_paragraph(doc, "Las contrasenas no deben almacenarse en texto plano. El sistema debe aplicar algoritmos de hash seguros como bcrypt o argon2, con sal unica por usuario. La base de datos no debe contener secretos directamente legibles.")

    doc.add_heading("Politica para el servidor Fedora", level=1)
    for item in [
        "Usar cuentas individuales para administracion.",
        "Evitar el uso de root para tareas comunes.",
        "Proteger el acceso SSH si se habilita durante la practica.",
        "Cerrar el puerto 5000 al finalizar la demostracion.",
        "Registrar cambios relevantes en el servidor y en el repositorio GitHub.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("Excepciones", level=1)
    add_paragraph(doc, "Durante el laboratorio academico se permiten credenciales debiles documentadas, como admin/admin123, unicamente para demostrar vulnerabilidades y justificar controles. Estas credenciales no deben utilizarse en entornos reales.")

    doc.add_heading("Sanciones o medidas correctivas", level=1)
    add_paragraph(doc, "El incumplimiento de esta politica debe corregirse mediante cambio inmediato de contrasena, revocacion temporal de acceso, revision de logs y actualizacion del control vulnerado.")

    doc.add_heading("Conclusiones", level=1)
    add_paragraph(doc, "Una politica de contrasenas reduce el riesgo de accesos no autorizados y fortalece la autenticacion. Para que sea efectiva, debe complementarse con implementacion tecnica, monitoreo y capacitacion de usuarios.")

    doc.add_heading("Bibliografia", level=1)
    for source in [SOURCES[3], SOURCES[5]]:
        add_paragraph(doc, source)

    polish_document(doc)
    path = OUT / "Politica_Uso_Contrasenas.docx"
    doc.save(path)
    return path


def build_controls_doc():
    doc = Document()
    set_doc_style(doc)
    add_cover(doc, "Establecimiento de Controles de Seguridad", PROJECT_TITLE)

    doc.add_heading("Objetivo", level=1)
    add_paragraph(doc, "Definir controles simples y aplicables para el uso seguro del sistema web de ventas y del servidor Fedora 39, con enfasis en contrasenas, acceso, operacion y cierre del laboratorio.")

    doc.add_heading("Controles para el sistema web", level=1)
    controls = [
        ["Autenticacion", "Uso de login obligatorio para acceder a productos, clientes, ventas y reportes."],
        ["Contrasenas", "Aplicar longitud minima, complejidad, bloqueo por intentos y hash seguro."],
        ["Sesiones", "Cerrar sesion al finalizar el uso del sistema."],
        ["Validacion de entradas", "Validar y sanitizar campos de productos, clientes y busqueda."],
        ["Autorizacion", "Definir roles como Administrador y Vendedor para limitar operaciones."],
        ["Registros", "Registrar inicios de sesion, errores y operaciones relevantes."],
        ["Backups", "Respaldar la base SQLite antes y despues de pruebas importantes."],
        ["Cookies seguras", "Configurar expiracion, HttpOnly, SameSite y Secure cuando exista HTTPS."],
    ]
    table = doc.add_table(rows=1, cols=2)
    set_cell_text(table.rows[0].cells[0], "Control", True)
    set_cell_text(table.rows[0].cells[1], "Descripcion", True)
    for row in controls:
        cells = table.add_row().cells
        set_cell_text(cells[0], row[0])
        set_cell_text(cells[1], row[1])
    style_table(table)

    doc.add_heading("Controles para el servidor Fedora", level=1)
    for item in [
        "Actualizar paquetes antes de la demostracion.",
        "Permitir el puerto 5000/tcp solo durante la presentacion.",
        "Verificar servicios activos con ss -tulnp.",
        "Listar reglas de firewall con sudo firewall-cmd --list-all.",
        "Revisar si el puerto 5355/LLMNR es necesario; si no lo es, deshabilitarlo para reducir superficie de ataque.",
        "Filtrar respuestas ICMP timestamp si no se requieren para la practica.",
        "Usar usuario no root para ejecutar la aplicacion.",
        "Apagar la VM o cerrar el servicio Flask al terminar la exposicion.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("Manual operativo simple", level=1)
    steps = [
        "Iniciar Fedora 39 en VirtualBox.",
        "Abrir terminal y entrar al repositorio del sistema.",
        "Activar el entorno virtual con source .venv/bin/activate.",
        "Ejecutar python app.py.",
        "Abrir el navegador en http://IP_DEL_SERVIDOR:5000.",
        "Ingresar con usuario admin y contrasena admin123 para la demostracion.",
        "Mostrar productos, clientes, ventas, reportes y busqueda.",
        "Ejecutar Nmap y Nessus desde la maquina de auditoria.",
        "Guardar capturas y reportes.",
        "Detener el servidor Flask y cerrar la VM.",
    ]
    for step in steps:
        add_number(doc, step)

    doc.add_heading("Controles posteriores recomendados", level=1)
    for item in [
        "Cambiar credenciales de prueba por credenciales robustas.",
        "Eliminar vulnerabilidades intencionales si el sistema se usa fuera del laboratorio.",
        "Desactivar debug=True.",
        "Reemplazar SQLite por una base administrada si el sistema escala a multiples usuarios.",
        "Implementar HTTPS si se expone fuera de red local.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("Conclusiones", level=1)
    add_paragraph(doc, "Los controles establecidos permiten reducir riesgos basicos del sistema web y del servidor, ademas de organizar la demostracion de auditoria como un laboratorio temporal, controlado y documentado.")

    polish_document(doc)
    path = OUT / "Establecimiento_Controles.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    for built in [build_main_report(), build_password_policy(), build_controls_doc()]:
        print(built)
